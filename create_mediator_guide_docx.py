import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("EMR-DHIS2 Interoperability Mediator")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66) # Navy Blue
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_p.add_run("Comprehensive Technical Specification, Deployment Architecture & Operational Guide")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    # Meta Box Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Project Domain:", "Ethiopian Ministry of Health (MOH) IMNID Ecosystem"),
        ("Service Name:", "EMR-DHIS2 Interoperability Mediator (v1.0.0)"),
        ("Supported Systems:", "OpenMRS 2.x (Bahmni) & OpenMRS 3.x (O3) ➔ DHIS2 Tracker"),
        ("Document Purpose:", "System Architecture, Gateway Integration, Network Topology & API Guide")
    ]
    for i, (label, val) in enumerate(meta_data):
        cell_lbl, cell_val = meta_table.rows[i].cells
        cell_lbl.text = label
        cell_lbl.paragraphs[0].runs[0].font.bold = True
        cell_lbl.paragraphs[0].runs[0].font.size = Pt(10)
        cell_val.text = val
        cell_val.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_background(cell_lbl, "F0F4F8")
        set_cell_background(cell_val, "F9FAFC")
        set_cell_margins(cell_lbl, top=80, bottom=80, left=120, right=120)
        set_cell_margins(cell_val, top=80, bottom=80, left=120, right=120)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
        return h

    # Section 1
    add_heading_1("1. Executive Summary & Overview")
    p = doc.add_paragraph()
    p.add_run("The ").font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    p.add_run("EMR-DHIS2 Interoperability Mediator").bold = True
    p.add_run(" is a lightweight, high-performance REST microservice built using FastAPI. It transitions the EMR-to-DHIS2 pipeline from a manual desktop utility into an automated, server-ready service designed for continuous operation within the Ethiopian Ministry of Health (MOH) digital health framework.\n\n")
    p.add_run("The mediator automates the complete end-to-end data pipeline in a single execution (~1.2 seconds):")
    
    bullets = [
        ("EMR Patient Extraction: ", "Connects via REST API to OpenMRS (Bahmni 2.x or OpenMRS 3.x) to query active patient visits."),
        ("Demographic Eligibility Filtering: ", "Enforces strict national MOH demographic rules (Maternal: Female & Age ≥ 10; Neonatal: Age ≤ 0)."),
        ("Form Schema & Concept Mapping: ", "Maps OpenMRS form schemas and observations directly to DHIS2 Data Elements and Attributes."),
        ("DHIS2 Tracker API Import: ", "Resolves facility Organisation Unit codes and posts Tracked Entity Instances, Enrollments, and Stage Events.")
    ]
    for b_title, b_desc in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r1 = bp.add_run(b_title)
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        bp.add_run(b_desc)

    # Section 2
    add_heading_1("2. OpenHIE Alignment & OpenHIM Gateway Integration")
    p = doc.add_paragraph()
    p.add_run("In the OpenHIE (Open Health Information Exchange) framework, interoperability architecture is divided into two distinct components: ")
    p.add_run("Interoperability Layer Gateways").bold = True
    p.add_run(" and ")
    p.add_run("Domain Mediator Microservices").bold = True
    p.add_run(".")
    
    # Table comparison
    add_heading_2("Component Comparison: OpenHIM Gateway vs. Our Mediator")
    comp_table = doc.add_table(rows=4, cols=3)
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Feature / Capability", "OpenHIM Gateway", "Our Mediator (mediator.py)"]
    for j, h in enumerate(headers):
        cell = comp_table.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "003366")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        
    rows_data = [
        ("Primary Role", "Central Security, Authentication & Audit Logging", "Clinical Data Extraction, Filtering & DHIS2 Import"),
        ("Domain Logic", "Protocol Routing & Pass-through Gateway", "Deep EMR Concept Mapping & Demographic Rules"),
        ("Execution Engine", "Node.js Reverse Proxy Gateway", "FastAPI / Uvicorn Python Microservice Engine")
    ]
    for i, row in enumerate(rows_data):
        r_cells = comp_table.rows[i+1].cells
        for j, val in enumerate(row):
            r_cells[j].text = val
            r_cells[j].paragraphs[0].runs[0].font.size = Pt(10)
            set_cell_background(r_cells[j], "F9FAFC" if i % 2 == 0 else "FFFFFF")
            set_cell_margins(r_cells[j], top=80, bottom=80, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    p_arch = doc.add_paragraph()
    p_arch.add_run("Production Architecture Flow:\n").bold = True
    p_arch.add_run("OpenMRS EMR ➔ OpenHIM Gateway (Security & Audit) ➔ Mediator Engine (mediator.py) ➔ DHIS2 Tracker API.")

    # Section 3
    add_heading_1("3. Pipeline Capabilities & Data Coverage")
    p_data = doc.add_paragraph()
    p_data.add_run("The mediator handles ").bold = True
    p_data.add_run("FULL CLINICAL DATA").bold = True
    p_data.add_run(" across all maternal and neonatal inpatient programs. It is not limited to patient demographics.")

    add_heading_2("Data Coverage Breakdown")
    data_bullets = [
        ("Demographic Attributes: ", "Record ID, MRN, First Name, Last Name, Age, Sex, Address (Zone, Wereda, Kebele), Phone Number."),
        ("Program Enrollments: ", "Enrolls entities in Maternal Inpatient Data (MID) or Neonatal Care Form (NCF) programs."),
        ("Clinical Stage Events: ", "Processes all repeatable stage forms including NICU Admission Careform, Progress Care Form for Referral Hospitals, Nurse Followup Sheet, KMC Ward Followup Sheet, Neonatal Referral Form, Discharge Careform, and Investigation Sheet."),
        ("Facility OrgUnit Resolution: ", "Automatically resolves facility codes (e.g. '1001' or MFL codes) to DHIS2 Organisation Unit UIDs.")
    ]
    for b_title, b_desc in data_bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r1 = bp.add_run(b_title)
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
        bp.add_run(b_desc)

    # Section 4
    add_heading_1("4. Deployment Strategies & Network Topology")
    p_dep = doc.add_paragraph()
    p_dep.add_run("Facility network configurations dictate how the mediator is deployed across health facilities in Ethiopia:")

    add_heading_2("Scenario A: Online Facilities (Hospital LAN with Internet)")
    p_on = doc.add_paragraph()
    p_on.add_run("When OpenMRS is accessible on a local hospital LAN (e.g., http://192.168.1.100:8080/openmrs) and the hospital has internet access:\n")
    p_on.add_run("1. The Mediator is installed inside the facility network on a local server.\n")
    p_on.add_run("2. It reads patient data locally over the hospital LAN.\n")
    p_on.add_run("3. It sends outbound HTTPS requests to central DHIS2 (https://imnid.mohdigitalhealth.gov.et).\n")
    p_on.add_run("Outbound connections are allowed by hospital firewalls even when inbound internet access is blocked.").italic = True

    add_heading_2("Scenario B: Offline Facilities (Disconnected LAN)")
    p_off = doc.add_paragraph()
    p_off.add_run("For health facilities without internet access to DHIS2:\n")
    p_off.add_run("1. Staff use the Desktop App (or Mediator locally) to extract and transform local EMR data to CSV.\n")
    p_off.add_run("2. The transformed CSV file is saved onto a USB flash drive.\n")
    p_off.add_run("3. The CSV is transported to a regional health hub or uploaded when internet connection is available.")

    # Section 5
    add_heading_1("5. API Specification & How to Use")
    p_api = doc.add_paragraph()
    p_api.add_run("The mediator exposes standard REST endpoints for automated scheduling and manual testing.")

    add_heading_2("Server Startup")
    doc.add_paragraph("To launch the mediator server locally or on a facility server:\n  python mediator.py").runs[0].font.name = 'Consolas'

    add_heading_2("API Endpoints")
    ep_table = doc.add_table(rows=4, cols=3)
    ep_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ep_headers = ["Method & Endpoint", "Purpose", "Target EMR Version"]
    for j, h in enumerate(ep_headers):
        cell = ep_table.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "003366")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)

    ep_data = [
        ("GET /health", "Health Check & Uptime Monitor", "All Systems"),
        ("POST /api/v1/sync/bahmni", "Full Sync Pipeline for Bahmni", "OpenMRS 2.x / Bahmni"),
        ("POST /api/v1/sync/o3", "Full Sync Pipeline with Form Schemas", "OpenMRS 3.x (O3)")
    ]
    for i, row in enumerate(ep_data):
        r_cells = ep_table.rows[i+1].cells
        for j, val in enumerate(row):
            r_cells[j].text = val
            r_cells[j].paragraphs[0].runs[0].font.size = Pt(10)
            set_cell_background(r_cells[j], "F9FAFC" if i % 2 == 0 else "FFFFFF")
            set_cell_margins(r_cells[j], top=80, bottom=80, left=100, right=100)

    add_heading_2("Sample Request Body (JSON)")
    json_text = """{
  "emr_base_url": "http://192.168.1.100:8080/openmrs",
  "emr_username": "superman",
  "emr_password": "Admin123",
  "facility_code": "1001",
  "visit_type_name": "Delivery",
  "start_date": "2026-01-01",
  "end_date": "2026-08-30",
  "dhis2_url": "https://imnid.mohdigitalhealth.gov.et",
  "dhis2_username": "admin",
  "dhis2_password": "district"
}"""
    jp = doc.add_paragraph()
    jrun = jp.add_run(json_text)
    jrun.font.name = 'Consolas'
    jrun.font.size = Pt(9.5)

    # Section 6
    add_heading_1("6. Desktop App vs. Mediator Hybrid Architecture")
    p_hy = doc.add_paragraph()
    p_hy.add_run("The platform is built as a ").bold = True
    p_hy.add_run("Hybrid System").bold = True
    p_hy.add_run(" offering two operating modes powered by the same underlying Python engine:\n\n")
    p_hy.add_run("1. Desktop GUI (main.py): ").bold = True
    p_hy.add_run("Used by hospital data clerks for visual date selection, file export, and manual syncs.\n")
    p_hy.add_run("2. REST Mediator (mediator.py): ").bold = True
    p_hy.add_run("Used by server background cron jobs and OpenHIM gateways for automated 24/7 synchronization.")

    output_path = "EMR_DHIS2_Interoperability_Mediator_Specification_Guide.docx"
    doc.save(output_path)
    print(f"Word Document successfully saved to {output_path}")

if __name__ == "__main__":
    create_document()
